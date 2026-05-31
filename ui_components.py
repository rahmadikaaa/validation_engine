"""
ui_components.py — Reusable UI components for SIPVAL Dashboard.
Provides HTML-rendered widgets and DOCX justification export.
"""

import streamlit as st
import pandas as pd
import io
from datetime import datetime
from typing import Dict, Any, List, Optional

from mock_data import format_currency, format_currency_full


# ============================================================
# Status Badge
# ============================================================

def render_status_badge(status: str) -> str:
    """Return HTML for a colored status badge."""
    status_upper = str(status).upper()
    if status_upper in ("VALID", "MATCH"):
        return '<span class="status-badge valid">✅ Valid</span>'
    elif status_upper == "PARTIAL_MATCH":
        return '<span class="status-badge partial">⚠️ Partial Match</span>'
    elif status_upper in ("INVALID", "NO_CANDIDATE"):
        return '<span class="status-badge invalid">❌ Invalid</span>'
    else:
        return f'<span class="status-badge neutral">◽ {status}</span>'


def render_status_text(status: str) -> str:
    """Return text label for status (used in non-HTML contexts)."""
    status_upper = str(status).upper()
    if status_upper in ("VALID", "MATCH"):
        return "✅ Valid"
    elif status_upper == "PARTIAL_MATCH":
        return "⚠️ Partial Match"
    elif status_upper in ("INVALID", "NO_CANDIDATE"):
        return "❌ Invalid"
    return f"◽ {status}"


# ============================================================
# KPI Card
# ============================================================

def render_kpi_card(icon: str, label: str, value: str, subtitle: str = "", color: str = "teal") -> str:
    """Return HTML for a glassmorphism KPI card."""
    return f"""
    <div class="kpi-card {color} fade-in">
        <div class="kpi-icon">{icon}</div>
        <div class="kpi-label">{label}</div>
        <div class="kpi-value">{value}</div>
        <div class="kpi-sub">{subtitle}</div>
    </div>
    """


# ============================================================
# Score Bar
# ============================================================

def render_score_bar(score: float, label: str = "Match Score") -> str:
    """Return HTML for a visual score progress bar.
    
    Score should be 0.0 to 1.0.
    """
    pct = max(0, min(100, round(score * 100)))
    if pct >= 70:
        level = "high"
    elif pct >= 40:
        level = "medium"
    else:
        level = "low"

    return f"""
    <div style="margin: 8px 0;">
        <div style="display: flex; justify-content: space-between; margin-bottom: 4px;">
            <span style="color: #94A3B8; font-size: 0.75rem; font-weight: 500;">{label}</span>
            <span style="color: #F1F5F9; font-size: 0.85rem; font-weight: 700;">{pct}%</span>
        </div>
        <div class="score-bar-container">
            <div class="score-bar-fill {level}" style="width: {pct}%;"></div>
        </div>
    </div>
    """


# ============================================================
# Page Header
# ============================================================

def render_page_header(icon: str, title: str, subtitle: str = ""):
    """Render a branded page header."""
    sub_html = f'<div class="subtitle">{subtitle}</div>' if subtitle else ""
    st.markdown(f"""
    <div class="sipval-header fade-in">
        <h1>{icon} {title}</h1>
        {sub_html}
    </div>
    """, unsafe_allow_html=True)


# ============================================================
# AI Reasoning Display
# ============================================================

REASONING_TYPE_LABELS = {
    "dimension_match": ("📐", "Kecocokan Dimensi"),
    "function_match": ("⚙️", "Kecocokan Fungsi"),
    "material_equivalence": ("🔩", "Kesetaraan Material"),
    "industrial_vs_consumer_grade": ("🏭", "Grade Industrial"),
    "quantity_or_unit": ("📦", "Kecocokan Satuan"),
}


def render_reasoning_cards(reasoning_list: List[Dict[str, Any]]):
    """Render AI reasoning as styled cards."""
    if not reasoning_list:
        st.info("Tidak ada data reasoning tersedia.")
        return

    for item in reasoning_list:
        r_type = item.get("type", "unknown")
        r_status = item.get("status", "unknown")
        r_notes = item.get("notes", "")

        icon, label = REASONING_TYPE_LABELS.get(r_type, ("🔍", r_type.replace("_", " ").title()))

        st.markdown(f"""
        <div class="reasoning-item">
            <div class="ri-header">
                <span class="ri-icon">{icon}</span>
                <span class="ri-type">{label}</span>
                <span class="ri-status {r_status}">{r_status.upper()}</span>
            </div>
            <div class="ri-notes">{r_notes}</div>
        </div>
        """, unsafe_allow_html=True)


# ============================================================
# Candidate Card
# ============================================================

def render_candidate_card(candidate: Dict[str, Any]):
    """Render the best candidate comparison card."""
    nama = candidate.get("comparison_product") or candidate.get("nama_produk", "N/A")
    vendor = candidate.get("vendor", "N/A")
    harga = candidate.get("comparison_price") or candidate.get("harga", 0)
    spec = candidate.get("comparison_specification") or candidate.get("spesifikasi", "N/A")
    satuan = candidate.get("satuan", "-")

    st.markdown(f"""
    <div class="candidate-card fade-in">
        <div class="cc-title">🏆 Kandidat Pembanding Terbaik</div>
        <div class="cc-name">{nama}</div>
        <div class="cc-details">
            <div class="cc-field">
                <span class="cc-field-label">Vendor</span>
                <span class="cc-field-value">{vendor}</span>
            </div>
            <div class="cc-field">
                <span class="cc-field-label">Harga</span>
                <span class="cc-field-value">{format_currency_full(harga)}</span>
            </div>
            <div class="cc-field">
                <span class="cc-field-label">Spesifikasi</span>
                <span class="cc-field-value">{spec}</span>
            </div>
            <div class="cc-field">
                <span class="cc-field-label">Satuan</span>
                <span class="cc-field-value">{satuan}</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)


# ============================================================
# Vendor Card
# ============================================================

def render_vendor_card(vendor: Dict[str, Any]):
    """Render a vendor summary card."""
    nama = vendor.get("nama", "Unknown")
    lokasi = vendor.get("lokasi", "-")
    items = vendor.get("items_supplied", 0)
    total = vendor.get("total_nilai", 0)
    reliability = vendor.get("reliability", 0)
    npwp = vendor.get("npwp", "-")
    valid = vendor.get("valid_count", 0)
    partial = vendor.get("partial_count", 0)
    invalid = vendor.get("invalid_count", 0)

    # Reliability color
    if reliability >= 90:
        rel_color = "#10B981"
    elif reliability >= 80:
        rel_color = "#F59E0B"
    else:
        rel_color = "#EF4444"

    st.markdown(f"""
    <div class="vendor-card fade-in">
        <div class="vc-name">🏢 {nama}</div>
        <div class="vc-location">📍 {lokasi} &nbsp;|&nbsp; NPWP: {npwp}</div>
        <div class="vc-stats">
            <div class="vc-stat">
                <span class="vc-stat-label">Item Dipasok</span>
                <span class="vc-stat-value">{items}</span>
            </div>
            <div class="vc-stat">
                <span class="vc-stat-label">Total Nilai</span>
                <span class="vc-stat-value">{format_currency(total)}</span>
            </div>
            <div class="vc-stat">
                <span class="vc-stat-label">Reliabilitas</span>
                <span class="vc-stat-value" style="color: {rel_color};">{reliability}%</span>
            </div>
            <div class="vc-stat">
                <span class="vc-stat-label">Status Validasi</span>
                <span class="vc-stat-value" style="font-size: 0.8rem;">✅{valid} ⚠️{partial} ❌{invalid}</span>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)


# ============================================================
# Info Card Row
# ============================================================

def render_info_card(label: str, value: str):
    """Render a simple info card."""
    st.markdown(f"""
    <div class="info-card">
        <div class="card-label">{label}</div>
        <div class="card-value">{value}</div>
    </div>
    """, unsafe_allow_html=True)


# ============================================================
# System Status
# ============================================================

def render_system_status():
    """Render system status indicators for sidebar."""
    return """
    <div style="padding: 8px 0;">
        <div style="color: #94A3B8; font-size: 0.7rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.05em; margin-bottom: 8px;">Status Sistem</div>
        <div style="margin-bottom: 6px;">
            <span class="sys-status online"><span class="dot"></span> Internal DB</span>
        </div>
        <div style="margin-bottom: 6px;">
            <span class="sys-status online"><span class="dot"></span> Reasoning Engine</span>
        </div>
        <div style="margin-bottom: 6px;">
            <span class="sys-status offline"><span class="dot"></span> INAPROC API</span>
        </div>
    </div>
    """


# ============================================================
# Justification Text Generator
# ============================================================

def generate_justification_text(
    item_data: Dict[str, Any],
    result_data: Dict[str, Any],
) -> str:
    """Generate a plain-text justification document."""
    nama = item_data.get("nama_barang") or item_data.get("Nama Barang dan Spesifikasi", "N/A")
    id_val = item_data.get("id_item") or item_data.get("ID", "N/A")
    kategori = item_data.get("Kategori") or item_data.get("kategori", "N/A")
    satuan = item_data.get("Satuan") or item_data.get("satuan", "N/A")
    harga_ref = item_data.get("Harga Satuan") or item_data.get("harga_satuan", "N/A")

    status = result_data.get("status", "N/A")
    score = result_data.get("score", 0)
    comp_product = result_data.get("comparison_product", "N/A")
    comp_spec = result_data.get("comparison_specification", "N/A")
    vendor = result_data.get("vendor", "N/A")
    comp_price = result_data.get("comparison_price", 0)
    deviation = result_data.get("deviation_notes", "")

    # Determine recommendation
    if status in ("VALID", "MATCH"):
        rekomendasi = "LAYAK — Item sesuai dengan referensi standar internal. Harga dan spesifikasi wajar, dapat diproses lebih lanjut."
    elif status == "PARTIAL_MATCH":
        rekomendasi = "PERLU REVIEW — Item memiliki kemiripan parsial dengan referensi. Terdapat deviasi minor yang perlu direview."
    else:
        rekomendasi = "TIDAK LAYAK — Item tidak memiliki referensi internal yang cocok atau deviasi terlalu tinggi. Wajib review manual."

    # Calculate deviation percentage
    try:
        harga_ref_val = float(harga_ref)
        comp_price_val = float(comp_price)
        if harga_ref_val > 0:
            dev_pct = ((comp_price_val - harga_ref_val) / harga_ref_val) * 100
            dev_str = f"{dev_pct:+.1f}%"
        else:
            dev_str = "N/A"
    except (ValueError, TypeError):
        dev_str = "N/A"

    now = datetime.now().strftime("%d %B %Y, %H:%M WIB")

    text = f"""
═══════════════════════════════════════════════════════════
         JUSTIFIKASI PENGADAAN BARANG/JASA
         SIPVAL — Sistem Validasi Pengadaan
═══════════════════════════════════════════════════════════

Tanggal     : {now}
Dokumen     : Justifikasi Validasi Pembanding

═══════════════════════════════════════════════════════════
1. INFORMASI ITEM PENGADAAN
═══════════════════════════════════════════════════════════

ID Item        : {id_val}
Nama Barang    : {nama}
Kategori       : {kategori}
Satuan         : {satuan}
Harga Referensi: {format_currency_full(harga_ref)}

═══════════════════════════════════════════════════════════
2. KANDIDAT PEMBANDING
═══════════════════════════════════════════════════════════

Produk Pembanding  : {comp_product}
Spesifikasi        : {comp_spec}
Vendor             : {vendor}
Harga Pembanding   : {format_currency_full(comp_price)}
Deviasi Harga      : {dev_str}

═══════════════════════════════════════════════════════════
3. HASIL VALIDASI AI
═══════════════════════════════════════════════════════════

Status Validasi    : {status}
Skor Kecocokan     : {score}
Catatan Deviasi    : {deviation}

═══════════════════════════════════════════════════════════
4. REKOMENDASI
═══════════════════════════════════════════════════════════

{rekomendasi}

═══════════════════════════════════════════════════════════
Dokumen ini dihasilkan secara otomatis oleh SIPVAL.
Sistem Validasi Pengadaan berbasis AI — v1.0 MVP
═══════════════════════════════════════════════════════════
"""
    return text.strip()


# ============================================================
# DOCX Justification Export
# ============================================================

def generate_justification_docx(
    item_data: Dict[str, Any],
    result_data: Dict[str, Any],
) -> Optional[bytes]:
    """Generate a .docx justification document. Returns bytes or None if python-docx unavailable."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return None

    nama = item_data.get("nama_barang") or item_data.get("Nama Barang dan Spesifikasi", "N/A")
    id_val = str(item_data.get("id_item") or item_data.get("ID", "N/A"))
    kategori = str(item_data.get("Kategori") or item_data.get("kategori", "N/A"))
    satuan = str(item_data.get("Satuan") or item_data.get("satuan", "N/A"))
    harga_ref = item_data.get("Harga Satuan") or item_data.get("harga_satuan", "N/A")

    status = result_data.get("status", "N/A")
    score = result_data.get("score", 0)
    comp_product = str(result_data.get("comparison_product", "N/A"))
    comp_spec = str(result_data.get("comparison_specification", "N/A"))
    vendor = str(result_data.get("vendor", "N/A"))
    comp_price = result_data.get("comparison_price", 0)
    deviation = str(result_data.get("deviation_notes", ""))

    # Recommendation
    if status in ("VALID", "MATCH"):
        rekomendasi = "LAYAK — Item sesuai dengan referensi standar internal. Harga dan spesifikasi wajar."
    elif status == "PARTIAL_MATCH":
        rekomendasi = "PERLU REVIEW — Item memiliki kemiripan parsial. Terdapat deviasi yang perlu diperiksa."
    else:
        rekomendasi = "TIDAK LAYAK — Referensi internal tidak cocok atau deviasi terlalu tinggi."

    now = datetime.now().strftime("%d %B %Y")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # Title
    title = doc.add_heading("JUSTIFIKASI PENGADAAN BARANG/JASA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(13, 148, 136)

    subtitle = doc.add_paragraph("SIPVAL — Sistem Validasi Pengadaan")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph(f"Tanggal: {now}")
    doc.add_paragraph("")

    # Section 1: Item Info
    doc.add_heading("1. Informasi Item Pengadaan", level=1)
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = "Light List Accent 1"
    fields1 = [
        ("ID Item", id_val),
        ("Nama Barang", str(nama)),
        ("Kategori", kategori),
        ("Satuan", satuan),
        ("Harga Referensi", format_currency_full(harga_ref)),
    ]
    for i, (label, val) in enumerate(fields1):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = val

    doc.add_paragraph("")

    # Section 2: Candidate
    doc.add_heading("2. Kandidat Pembanding", level=1)
    table2 = doc.add_table(rows=4, cols=2)
    table2.style = "Light List Accent 1"
    fields2 = [
        ("Produk Pembanding", comp_product),
        ("Spesifikasi", comp_spec),
        ("Vendor", vendor),
        ("Harga Pembanding", format_currency_full(comp_price)),
    ]
    for i, (label, val) in enumerate(fields2):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = val

    doc.add_paragraph("")

    # Section 3: Validation
    doc.add_heading("3. Hasil Validasi AI", level=1)
    table3 = doc.add_table(rows=3, cols=2)
    table3.style = "Light List Accent 1"
    fields3 = [
        ("Status Validasi", str(status)),
        ("Skor Kecocokan", str(score)),
        ("Catatan Deviasi", deviation),
    ]
    for i, (label, val) in enumerate(fields3):
        table3.rows[i].cells[0].text = label
        table3.rows[i].cells[1].text = val

    doc.add_paragraph("")

    # Section 4: Recommendation
    doc.add_heading("4. Rekomendasi", level=1)
    rec_para = doc.add_paragraph(rekomendasi)
    for run in rec_para.runs:
        run.font.bold = True
        run.font.size = Pt(11)

    doc.add_paragraph("")

    # Footer
    footer = doc.add_paragraph("Dokumen ini dihasilkan secara otomatis oleh SIPVAL — Sistem Validasi Pengadaan v1.0")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(148, 163, 184)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# Spec Grid — Structured Specification Display
# ============================================================

def render_spec_grid(spec_text: str, item_dict: Dict[str, Any]) -> str:
    """Parse spec text into structured attribute grid HTML.

    Attempts to extract known attributes from the spec string.
    Falls back to line-by-line display if structured parsing fails.
    """
    import re

    # Known attribute patterns (case-insensitive)
    ATTR_PATTERNS = [
        ("Part Number",  r"(?:part\s*(?:number|no|num|#))[:\s]+([^\n;,]+)"),
        ("Brand / Merek", r"(?:brand|merek|merk)[:\s]+([^\n;,]+)"),
        ("Fungsi",       r"(?:function|fungsi|kegunaan)[:\s]+([^\n;,]+)"),
        ("Kapasitas",    r"(?:capacity|kapasitas|volume)[:\s]+([^\n;,]+)"),
        ("Dimensi",      r"(?:dimension|dimensi|ukuran|size)[:\s]+([^\n;,]+)"),
        ("Material",     r"(?:material|bahan)[:\s]+([^\n;,]+)"),
        ("Tegangan",     r"(?:voltage|tegangan|volt)[:\s]+([^\n;,]+)"),
        ("Daya",         r"(?:power|daya|watt)[:\s]+([^\n;,]+)"),
        ("Satuan",       r"(?:satuan|unit)[:\s]+([^\n;,]+)"),
        ("Kualitas",     r"(?:grade|kualitas|class)[:\s]+([^\n;,]+)"),
    ]

    # Get spec from multiple possible sources
    raw_spec = spec_text or ""
    if not raw_spec:
        raw_spec = str(item_dict.get("spesifikasi") or item_dict.get("Spesifikasi") or "")
    if not raw_spec:
        raw_spec = str(item_dict.get("nama_barang") or item_dict.get("Nama Barang dan Spesifikasi") or "")

    # Split spec string if it contains "Spesifikasi:"
    if "Spesifikasi:" in raw_spec:
        parts = raw_spec.split("Spesifikasi:", 1)
        raw_spec = parts[1].strip() if len(parts) > 1 else raw_spec

    # Try structured extraction
    attrs = []
    for label, pattern in ATTR_PATTERNS:
        match = re.search(pattern, raw_spec, re.IGNORECASE)
        if match:
            value = match.group(1).strip().rstrip(".,;")
            if value and len(value) < 120:
                attrs.append((label, value))

    # If we found no structured attributes, split by common delimiters
    if not attrs and raw_spec:
        # Try splitting by semicolons or newlines
        lines = re.split(r"[;\n]+", raw_spec)
        for line in lines[:6]:
            line = line.strip()
            if ":" in line:
                key, _, val = line.partition(":")
                key = key.strip()
                val = val.strip()
                if key and val and len(key) < 40 and len(val) < 120:
                    attrs.append((key, val))
            elif line and len(line) > 3:
                attrs.append(("Keterangan", line))

    # Fallback: show full spec text in a single card
    if not attrs:
        if raw_spec:
            truncated = raw_spec[:300] + ("..." if len(raw_spec) > 300 else "")
            return f"""
            <div class="spec-grid" style="grid-template-columns: 1fr;">
                <div class="spec-attr">
                    <div class="spec-attr-label">Spesifikasi</div>
                    <div class="spec-attr-value">{truncated}</div>
                </div>
            </div>"""
        else:
            return '<div style="color: #64748B; font-size: 0.82rem; font-style: italic;">Data spesifikasi tidak tersedia.</div>'

    # Render attribute grid
    items_html = ""
    for label, value in attrs[:9]:  # max 9 tiles (3×3 grid)
        items_html += f"""
        <div class="spec-attr">
            <div class="spec-attr-label">{label}</div>
            <div class="spec-attr-value">{value}</div>
        </div>"""

    return f'<div class="spec-grid">{items_html}</div>'


# ============================================================
# Validation Checklist
# ============================================================

# Maps reasoning_type → display label for checklist
CHECKLIST_CRITERIA = [
    ("function_match",             "⚙️", "Function Match"),
    ("dimension_match",            "📐", "Size / Dimension Match"),
    ("material_equivalence",       "🔩", "Material Match"),
    ("industrial_vs_consumer_grade","🏭", "Brand / Grade Equivalence"),
    ("quantity_or_unit",           "📦", "Unit Match"),
    ("vendor_location",            "📍", "Vendor Location Match"),
]


def _map_reasoning_status(r_status: str) -> str:
    """Normalize reasoning status to checklist class."""
    s = str(r_status).lower()
    if s in ("match", "pass", "valid", "equivalent"):
        return "pass"
    elif s in ("partial", "review", "partial_match", "close"):
        return "review"
    elif s in ("mismatch", "fail", "invalid"):
        return "fail"
    return "unknown"


def render_validation_checklist(reasoning_list: List[Dict[str, Any]]) -> str:
    """Render the validation checklist from reasoning output as HTML.

    Maps reasoning types to the 6 standard procurement checklist criteria.
    """
    STATUS_ICON = {
        "pass":    ("✓", "Pass"),
        "review":  ("⚠", "Review"),
        "fail":    ("✗", "Fail"),
        "unknown": ("◽", "N/A"),
    }

    # Build a lookup from reasoning_type → status
    reasoning_lookup: Dict[str, str] = {}
    for r in reasoning_list:
        r_type = r.get("type", "")
        r_status = _map_reasoning_status(r.get("status", "unknown"))
        reasoning_lookup[r_type] = r_status

    # For vendor_location we derive from location-related notes if present
    if "vendor_location" not in reasoning_lookup:
        # Check notes of any reasoning item for location keywords
        for r in reasoning_list:
            notes = str(r.get("notes", "")).lower()
            if any(k in notes for k in ["jakarta", "jabodetabek", "lokasi", "location", "wilayah"]):
                reasoning_lookup["vendor_location"] = _map_reasoning_status(r.get("status", "unknown"))
                break

    items_html = ""
    for r_type, icon, label in CHECKLIST_CRITERIA:
        status = reasoning_lookup.get(r_type, "unknown")
        s_icon, s_text = STATUS_ICON[status]
        items_html += f"""
        <div class="checklist-item {status}">
            <span class="checklist-icon">{icon}</span>
            <span class="checklist-label">{label}</span>
            <span class="checklist-status-text">{s_icon} {s_text}</span>
        </div>"""

    return f'<div class="checklist-grid">{items_html}</div>'


# ============================================================
# Candidate Comparison Table
# ============================================================

# Source priority ordering + badge class
SOURCE_INFO = [
    ("INAPROC",         "inaproc",  1),
    ("Tokopedia",       "tokopedia", 2),
    ("Shopee",          "shopee",   3),
    ("Official Vendor", "vendor",   4),
    ("Vendor A",        "vendor",   4),
    ("Vendor B",        "vendor",   4),
    ("Vendor C",        "vendor",   4),
]


def _source_badge(vendor_str: str) -> tuple:
    """Return (badge_html, priority) for a vendor string."""
    vendor_lower = str(vendor_str).lower()
    for label, css_class, priority in SOURCE_INFO:
        if label.lower() in vendor_lower:
            return f'<span class="source-badge {css_class}">{label}</span>', priority
    # Default
    return f'<span class="source-badge vendor">{vendor_str[:16]}</span>', 5


def _location_badge(location_str: str) -> tuple:
    """Return (badge_html, priority) for a location string."""
    loc_lower = str(location_str).lower()
    if "jakarta" in loc_lower and "bodetabek" not in loc_lower:
        return '<span class="location-badge jakarta">Jakarta</span>', 1
    elif any(k in loc_lower for k in ["bogor", "depok", "tangerang", "bekasi", "jabodetabek"]):
        return '<span class="location-badge jabodetabek">Jabodetabek</span>', 2
    else:
        return f'<span class="location-badge national">{location_str[:14] or "National"}</span>', 3


def _calc_deviation(ref_price, cand_price) -> tuple:
    """Return (pct_str, css_class, float_value)."""
    try:
        ref = float(ref_price)
        cand = float(cand_price)
        if ref > 0:
            pct = ((cand - ref) / ref) * 100
            css_class = "positive" if pct > 0 else ("negative" if pct < 0 else "neutral")
            return f"{pct:+.1f}%", css_class, pct
    except (ValueError, TypeError):
        pass
    return "N/A", "neutral", 0.0


def render_candidate_comparison_table(
    candidates: List[Dict[str, Any]],
    reference_price: float,
    selected_candidates: set,
    rejected_candidates: set,
) -> str:
    """Render full candidate comparison table as HTML.

    Returns the HTML string for the table (buttons are handled by Streamlit separately).
    """
    if not candidates:
        return '<div style="color: #64748B; font-size: 0.82rem; font-style: italic; padding: 12px 0;">Tidak ada kandidat tersedia.</div>'

    # Sort by source priority, then location priority
    def sort_key(c):
        _, src_prio = _source_badge(c.get("vendor", ""))
        _, loc_prio = _location_badge(c.get("lokasi", ""))
        return (src_prio, loc_prio)

    sorted_cands = sorted(candidates, key=sort_key)

    header = """
    <div class="cand-table-header">
        <div class="cand-table-header-cell">Source</div>
        <div class="cand-table-header-cell">Vendor</div>
        <div class="cand-table-header-cell">Product</div>
        <div class="cand-table-header-cell">Location</div>
        <div class="cand-table-header-cell">Cand. Price</div>
        <div class="cand-table-header-cell">Deviation</div>
        <div class="cand-table-header-cell">Confidence</div>
        <div class="cand-table-header-cell">Val. Status</div>
        <div class="cand-table-header-cell">Actions</div>
    </div>"""

    rows_html = header
    for i, cand in enumerate(sorted_cands):
        vendor = str(cand.get("vendor", "N/A"))
        product = str(cand.get("nama_produk") or cand.get("comparison_product") or "N/A")
        price = cand.get("harga") or cand.get("comparison_price") or 0
        lokasi = str(cand.get("lokasi", "National"))
        score = cand.get("score", None)
        val_status = str(cand.get("status", "—"))

        src_badge_html, _ = _source_badge(vendor)
        loc_badge_html, _ = _location_badge(lokasi)
        dev_str, dev_class, dev_val = _calc_deviation(reference_price, price)

        # Confidence pill
        if score is not None:
            score_pct = round(float(score) * 100)
            pill_level = "high" if score_pct >= 70 else ("medium" if score_pct >= 40 else "low")
            conf_html = f'<span class="confidence-pill {pill_level}">{score_pct}%</span>'
        else:
            conf_html = '<span class="confidence-pill none">—</span>'

        # Row selection state
        cand_id = f"cand_{i}"
        row_class = "cand-table-row"
        if cand_id in selected_candidates:
            row_class += " selected-cand"
        elif cand_id in rejected_candidates:
            row_class += " rejected-cand"

        # Price
        try:
            price_str = f"Rp {float(price):,.0f}".replace(",", ".")
        except (ValueError, TypeError):
            price_str = "N/A"

        # Truncate product name
        product_truncated = product[:45] + ("..." if len(product) > 45 else "")

        # Status badge  
        status_map = {
            "VALID": ("✅", "#10B981"),
            "MATCH": ("✅", "#10B981"),
            "PARTIAL_MATCH": ("⚠️", "#F59E0B"),
            "INVALID": ("❌", "#EF4444"),
        }
        s_icon, s_color = status_map.get(val_status.upper(), ("◽", "#94A3B8"))

        rows_html += f"""
        <div class="{row_class}" id="cand-row-{i}">
            <div>{src_badge_html}</div>
            <div style="color: #94A3B8; font-size: 0.78rem; overflow: hidden; text-overflow: ellipsis; white-space: nowrap;">{vendor[:22]}</div>
            <div style="color: #CBD5E1; font-size: 0.8rem; line-height: 1.3;">{product_truncated}</div>
            <div>{loc_badge_html}</div>
            <div style="color: #FBBF24; font-size: 0.8rem; font-weight: 600;">{price_str}</div>
            <div class="deviation-cell {dev_class}">{dev_str}</div>
            <div>{conf_html}</div>
            <div style="color: {s_color}; font-size: 0.78rem; font-weight: 600;">{s_icon} {val_status}</div>
            <div style="color: #64748B; font-size: 0.72rem;">↓ see buttons</div>
        </div>"""

    return rows_html


# ============================================================
# Procurement Reasoning Block
# ============================================================

def render_procurement_reasoning_block(
    reasoning_list: List[Dict[str, Any]],
    status: str,
    score: float,
    deviation_notes: str = "",
) -> str:
    """Render the AI-generated procurement reasoning narrative as HTML."""
    status_upper = str(status).upper()

    if status_upper in ("VALID", "MATCH"):
        verdict_class = "valid"
        verdict_text = "✅ VALID"
        summary_text = (
            f"Spesifikasi fungsional sesuai dengan referensi internal. "
            f"Skor kecocokan {round(float(score) * 100) if score else 0}%. "
            f"{'Deviasi harga dapat diterima dengan justifikasi yang memadai.' if deviation_notes else 'Item dapat diproses untuk pengadaan.'}"
        )
    elif status_upper == "PARTIAL_MATCH":
        verdict_class = "partial"
        verdict_text = "⚠️ PARTIAL MATCH"
        summary_text = (
            f"Item memiliki kemiripan parsial dengan kandidat pembanding. "
            f"Skor kecocokan {round(float(score) * 100) if score else 0}%. "
            f"Perlu review lebih lanjut sebelum diproses."
        )
    else:
        verdict_class = "invalid"
        verdict_text = "❌ TIDAK VALID"
        summary_text = (
            f"Item tidak memiliki referensi pembanding yang cukup kuat. "
            f"Skor kecocokan {round(float(score) * 100) if score else 0}%. "
            f"Wajib dilakukan review manual dan justifikasi khusus."
        )

    # Build reasoning bullets
    bullets_html = ""
    REASONING_LABELS = {
        "dimension_match":              "Kecocokan dimensi diperiksa",
        "function_match":               "Fungsi utama diverifikasi",
        "material_equivalence":         "Kesetaraan material dievaluasi",
        "industrial_vs_consumer_grade": "Grade produk dibandingkan",
        "quantity_or_unit":             "Satuan pengukuran divalidasi",
    }

    if reasoning_list:
        for r in reasoning_list:
            r_type = r.get("type", "")
            r_status = r.get("status", "")
            r_notes = r.get("notes", "")
            label = REASONING_LABELS.get(r_type, r_type.replace("_", " ").title())
            status_label = {"match": "✓", "partial": "~", "mismatch": "✗"}.get(
                str(r_status).lower(), "○"
            )
            note_text = f": {r_notes}" if r_notes else ""
            bullets_html += f'<div class="reasoning-bullet">{status_label} {label}{note_text}</div>'
    else:
        bullets_html = '<div class="reasoning-bullet">Jalankan validasi untuk mendapatkan detail reasoning.</div>'

    # Deviation note
    dev_note_html = ""
    if deviation_notes:
        dev_note_html = f'<div style="color: #F59E0B; font-size: 0.78rem; margin-top: 10px; padding: 8px 12px; background: rgba(245,158,11,0.08); border-radius: 6px;">⚠️ Catatan Deviasi: {deviation_notes}</div>'

    return f"""
    <div class="reasoning-block">
        <div class="reasoning-verdict {verdict_class}">{verdict_text}</div>
        {bullets_html}
        {dev_note_html}
        <div class="reasoning-summary">{summary_text}</div>
    </div>"""


# ============================================================
# High Deviation Warning
# ============================================================

DEVIATION_WARNING_REASONS = [
    "Produk grade industrial (bukan consumer grade)",
    "Barang impor atau komponen khusus",
    "Distributor resmi / official vendor pricing",
    "Perbedaan kuantitas atau kemasan (satuan berbeda)",
    "Perbedaan lokasi pengiriman (luar Jabodetabek)",
    "Kelangkaan stok atau musim permintaan tinggi",
]


def render_deviation_warning(deviation_pct: float) -> str:
    """Render a high deviation warning card if deviation > 20%."""
    if abs(deviation_pct) <= 20:
        return ""

    direction = "lebih tinggi" if deviation_pct > 0 else "lebih rendah"
    reasons_html = "".join(
        f'<div class="deviation-warning-reason">{r}</div>'
        for r in DEVIATION_WARNING_REASONS
    )

    return f"""
    <div class="deviation-warning">
        <div class="deviation-warning-title">
            ⚠️ Deviasi Harga Melebihi Ambang Batas ({deviation_pct:+.1f}%)
        </div>
        <div style="color: #94A3B8; font-size: 0.8rem; margin-bottom: 10px;">
            Harga kandidat <strong style="color: #FBBF24;">{abs(deviation_pct):.1f}%</strong> {direction} dari harga referensi.
            Kemungkinan penyebab:
        </div>
        {reasons_html}
    </div>"""

